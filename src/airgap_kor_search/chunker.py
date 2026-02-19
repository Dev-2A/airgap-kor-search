"""문서 처리 & 청킹 모듈

다양한 포맷(txt, pdf, docx, hwp)의 문서를 읽고,
검색에 적합한 크기로 텍스트를 분할합니다.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

from airgap_kor_search.config import ChunkConfig

logger = logging.getLogger(__name__)


# ── 데이터 클래스 ─────────────────────────────────────────


@dataclass
class Document:
    """읽어들인 원본 문서"""

    path: str
    text: str
    metadata: dict = field(default_factory=dict)


@dataclass
class Chunk:
    """분할된 텍스트 조각"""

    text: str
    doc_path: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: dict = field(default_factory=dict)

    @property
    def chunk_id(self) -> str:
        """고유 식별자: 문서경로::청크인덱스"""
        return f"{self.doc_path}::{self.chunk_index}"


# ── 문서 리더 ─────────────────────────────────────────────


class DocumentReader:
    """파일 포맷별 텍스트 추출기

    지원 포맷: .txt, .md, .pdf, .docx
    """

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}

    @classmethod
    def read(cls, path: Path | str) -> Document:
        """파일을 읽어 Document 객체로 변환합니다.

        Args:
            path: 문서 파일 경로

        Returns:
            Document 객체

        Raises:
            FileNotFoundError: 파일이 존재하지 않을 때
            ValueError: 지원하지 않는 파일 형식일 때
        """
        path = Path(path)

        if not path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {path}")

        ext = path.suffix.lower()
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"지원하지 않는 파일 형식입니다: {ext}\n"
                f"지원형식: {', '.join(sorted(cls.SUPPORTED_EXTENSIONS))}"
            )

        reader_map = {
            ".txt": cls._read_text,
            ".md": cls._read_text,
            ".pdf": cls._read_pdf,
            ".docx": cls._read_docx,
        }

        text = reader_map[ext](path)
        logger.info("문서 읽기 완료: %s (%d자)", path.name, len(text))

        return Document(
            path=str(path),
            text=text,
            metadata={"filename": path.name, "extension": ext},
        )

    @classmethod
    def read_directory(
        cls,
        directory: Path | str,
        recursive: bool = True,
    ) -> list[Document]:
        """디렉토리 내 모든 지원 문서를 읽습니다.

        Args:
            directory: 문서가 들어있는 디렉토리 경로
            recursive: 하위 디렉토리도 탐색할지 여부

        Returns:
            Document 리스트
        """
        directory = Path(directory)
        if not directory.is_dir():
            raise NotADirectoryError(f"디렉토리가 아닙니다: {directory}")

        pattern = "**/*" if recursive else "*"
        documents = []

        for file_path in sorted(directory.glob(pattern)):
            if file_path.is_file() and file_path.suffix.lower() in cls.SUPPORTED_EXTENSIONS:
                try:
                    doc = cls.read(file_path)
                    documents.append(doc)
                except Exception as e:
                    logger.warning("문서 읽기 실패 (%s): %s", file_path, e)

        logger.info("총 %d개 문서 로드 완료 (%s)", len(documents), directory)
        return documents

    @staticmethod
    def _read_text(path: Path) -> str:
        """텍스트/마크다운 파일 읽기"""
        encodings = ["utf-8", "cp949", "euc-kr"]
        for enc in encodings:
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError(
            "multi", b"", 0, 1,
            f"지원하는 인코딩으로 읽을 수 없습니다: {path}"
        )

    @staticmethod
    def _read_pdf(path: Path) -> str:
        """PDF 파일 텍스트 추출"""
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    @staticmethod
    def _read_docx(path: Path) -> str:
        """DOCX 파일 텍스트 추출"""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)


# ── 텍스트 청킹 ──────────────────────────────────────────


class TextChunker:
    """텍스트를 검색에 적합한 크기로 분할합니다.

    한국어 문서에 적합하도록 문단/문장 경계를 존중하며 분할합니다.

    사용 예시:
        >>> chunker = TextChunker.from_config(chunk_config)
        >>> chunks = chunker.chunk_document(document)
    """

    # 분할 우선순위: 문단 → 문장 → 강제 분할
    PARAGRAPH_SEPARATORS = ["\n\n", "\n"]
    SENTENCE_ENDINGS = re.compile(r"(?<=[.!?。！？])\s+")

    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        min_chunk_length: int = 50,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_length = min_chunk_length

    @classmethod
    def from_config(cls, config: ChunkConfig) -> "TextChunker":
        """ChunkConfig에서 TextChunker를 생성합니다."""
        return cls(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            min_chunk_length=config.min_chunk_length,
        )

    def chunk_text(self, text: str) -> list[str]:
        """텍스트를 청크 문자열 리스트로 분할합니다.

        Args:
            text: 원본 텍스트

        Returns:
            분할된 텍스트 리스트
        """
        if not text or not text.strip():
            return []

        text = self._normalize(text)

        # 1단계: 문단 단위로 먼저 분리
        paragraphs = self._split_paragraphs(text)

        # 2단계: 큰 문단은 문장 단위로 추가 분할
        segments = []
        for para in paragraphs:
            if len(para) <= self.chunk_size:
                segments.append(para)
            else:
                segments.extend(self._split_sentences(para))

        # 3단계: 세그먼트들을 chunk_size에 맞게 병합/분할
        chunks = self._merge_segments(segments)

        # 4단계: 너무 짧은 청크 제거
        chunks = [c for c in chunks if len(c.strip()) >= self.min_chunk_length]

        return chunks

    def chunk_document(self, document: Document) -> list[Chunk]:
        """Document를 Chunk 리스트로 분할합니다.

        Args:
            document: 원본 문서

        Returns:
            Chunk 리스트
        """
        chunk_texts = self.chunk_text(document.text)
        chunks = []

        current_pos = 0
        for i, chunk_text in enumerate(chunk_texts):
            # 원본 텍스트에서의 위치 추적
            start = document.text.find(chunk_text[:50], current_pos)
            if start == -1:
                start = current_pos
            end = start + len(chunk_text)
            current_pos = max(current_pos, start + 1)

            chunks.append(
                Chunk(
                    text=chunk_text,
                    doc_path=document.path,
                    chunk_index=i,
                    start_char=start,
                    end_char=end,
                    metadata={**document.metadata, "chunk_total": len(chunk_texts)},
                )
            )

        logger.info(
            "문서 청킹 완료: %s → %d개 청크",
            document.metadata.get("filename", document.path),
            len(chunks),
        )
        return chunks

    def chunk_documents(self, documents: list[Document]) -> list[Chunk]:
        """여러 Document를 한꺼번에 청킹합니다."""
        all_chunks = []
        for doc in documents:
            all_chunks.extend(self.chunk_document(doc))
        logger.info("전체 청킹 완료: %d개 문서 → %d개 청크", len(documents), len(all_chunks))
        return all_chunks

    @staticmethod
    def _normalize(text: str) -> str:
        """텍스트 정규화 (연속 공백 제거 등)"""
        # 연속된 빈 줄을 2개로 통일
        text = re.sub(r"\n{3,}", "\n\n", text)
        # 탭을 공백으로
        text = text.replace("\t", " ")
        # 연속 공백을 하나로 (줄바꿈 제외)
        text = re.sub(r"[^\S\n]+", " ", text)
        return text.strip()

    @staticmethod
    def _split_paragraphs(text: str) -> list[str]:
        """문단 단위로 분리"""
        paragraphs = text.split("\n\n")
        return [p.strip() for p in paragraphs if p.strip()]

    @classmethod
    def _split_sentences(cls, text: str) -> list[str]:
        """문장 단위로 분리"""
        sentences = cls.SENTENCE_ENDINGS.split(text)
        return [s.strip() for s in sentences if s.strip()]

    def _merge_segments(self, segments: list[str]) -> list[str]:
        """세그먼트들을 chunk_size에 맞게 병합합니다.

        오버랩을 적용하여 청크 간 문맥 연속성을 유지합니다.
        """
        if not segments:
            return []

        chunks = []
        current = segments[0]

        for segment in segments[1:]:
            # 현재 청크에 다음 세그먼트를 추가해도 크기 이내인 경우
            combined = current + " " + segment
            if len(combined) <= self.chunk_size:
                current = combined
            else:
                # 현재 청크를 확정하고 새 청크 시작
                chunks.append(current)

                # 오버랩 적용: 이전 청크의 끝부분을 새 청크의 시작으로
                if self.chunk_overlap > 0 and len(current) > self.chunk_overlap:
                    overlap_text = current[-self.chunk_overlap :]
                    # 단어 경계에서 자르기
                    space_pos = overlap_text.find(" ")
                    if space_pos != -1:
                        overlap_text = overlap_text[space_pos + 1 :]
                    current = overlap_text + " " + segment
                else:
                    current = segment

                # 병합 후에도 chunk_size 초과 시 강제 분할
                while len(current) > self.chunk_size:
                    cut_point = self._find_cut_point(current, self.chunk_size)
                    chunks.append(current[:cut_point])
                    current = current[cut_point:].strip()

        # 마지막 청크
        if current.strip():
            chunks.append(current)

        return chunks

    @staticmethod
    def _find_cut_point(text: str, max_length: int) -> int:
        """강제 분할 시 최적의 분할 지점을 찾습니다.

        문장 끝 → 쉼표/세미콜론 → 공백 → max_length 순으로 탐색합니다.
        """
        search_text = text[:max_length]

        # 문장 끝 기호 찾기 (뒤에서부터)
        for pattern in [". ", "! ", "? ", "。", "！", "？"]:
            pos = search_text.rfind(pattern)
            if pos != -1 and pos > max_length // 2:
                return pos + len(pattern)

        # 쉼표/세미콜론
        for pattern in [", ", "; ", "， "]:
            pos = search_text.rfind(pattern)
            if pos != -1 and pos > max_length // 2:
                return pos + len(pattern)

        # 공백
        pos = search_text.rfind(" ")
        if pos != -1 and pos > max_length // 2:
            return pos + 1

        # 최후의 수단: 그냥 자르기
        return max_length
